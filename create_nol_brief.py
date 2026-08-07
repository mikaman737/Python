from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "NOL_통합예약_패키지_수요분석_기획서.docx"

BLUE = "2E74B5"
NAVY = "1F4D78"
GRAY = "5B6573"
LIGHT = "E8EEF5"

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")

def set_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run(r, 12, True, BLUE)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run(r, 10.5)
    return p

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.68)
section.left_margin = Inches(0.78)
section.right_margin = Inches(0.78)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.10

# Header
hp = section.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("PROJECT BRIEF | 2026")
set_run(hr, 8.5, True, GRAY)

# Title block
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("수도권 여행객의 여행 복잡도 기반")
set_run(r, 20, True, NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(7)
r = p.add_run("NOL 통합 예약 패키지 수요 분석 기획서")
set_run(r, 20, True, NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(7)
r = p.add_run("목적: 복잡한 여행을 먼저 식별해 숙소·교통·티켓·레저를 한 번에 예약할 통합 패키지의 우선 타깃을 제안한다.")
set_run(r, 10.5, False, GRAY)

add_heading(doc, "1. 추진 배경")
add_body(doc, "놀유니버스는 2026년 9월 NOL 인터파크투어와 NOL 티켓을 NOL로 통합하고, 이후 트리플도 순차 결합하겠다고 발표했다. 상품이 한 플랫폼에서 연결되는 만큼, NOL은 모든 고객이 아닌 ‘통합 예약의 효용이 큰 여행’을 우선 공략할 필요가 있다.")

src = doc.add_paragraph()
src.paragraph_format.space_after = Pt(4)
sr = src.add_run("근거: 놀유니버스 공식 뉴스룸(2026.06.12), ‘NOL 여행·여가·문화 플랫폼 하나로’")
set_run(sr, 8.5, False, GRAY)

add_heading(doc, "2. 분석 데이터 및 핵심 변수")
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = "Table Grid"
headers = ["분석 영역", "주요 컬럼", "분석 목적"]
for cell, text in zip(table.rows[0].cells, headers):
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    pr = cell.paragraphs[0]
    pr.paragraph_format.space_after = Pt(0)
    run = pr.add_run(text)
    set_run(run, 9.2, True, NAVY)
rows = [
    ("여행 구조", "TRAVEL_ID, 시작·종료일, 여행 목적", "여행기간 및 기본 여행유형 분류"),
    ("동행·이동", "동반자 정보, 출발지·도착지, 이동수단", "동행 규모와 이동 복잡도 산출"),
    ("방문·활동", "방문 순서, 방문지명, 시군구, 유형, 활동유형, 사진 수", "동선 다양성 및 여행 경험 밀도 산출"),
    ("여행 규모", "사전·숙박·이동·활동 소비의 결제일·금액", "총지출·사전준비 시점 보조 비교"),
]
for row in rows:
    cells = table.add_row().cells
    for cell, text in zip(cells, row):
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        rr = p.add_run(text)
        set_run(rr, 8.8)
set_table_widths(table, [1650, 4150, 3560])

add_heading(doc, "3. 핵심 분석 및 제안 방향")
for title, body in [
    ("여행 복잡도 점수", "여행기간, 동행자 수, 방문지·시군구 수, 이동 수·이동수단 수, 활동 수를 결합해 여행별 예약·준비 부담을 수치화한다."),
    ("고복잡도 여행 세그먼트", "점수를 기준으로 2박 이상 다지역 친구여행, 가족 체험형 등 실제 여행 유형을 도출하고, 규모·지출·사전결제 리드타임을 비교한다."),
    ("통합 패키지 우선순위", "고복잡도 세그먼트의 반복 동선과 방문지 유형을 확인해 숙소+교통, 티켓+숙소, 가족 레저+숙소 등 가장 적합한 통합 예약 방식을 제안한다."),
]:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.05)
    p.paragraph_format.first_line_indent = Inches(-0.05)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.06
    r = p.add_run(title + " - ")
    set_run(r, 10, True, NAVY)
    r = p.add_run(body)
    set_run(r, 10)

add_heading(doc, "4. 최종 산출물 및 한계")
add_body(doc, "최종 산출물은 ‘여행 복잡도별 세그먼트 × 반복 동선 × 통합 예약 패키지’ 매트릭스다. 구매처·예약처 정보가 불완전하므로 외부 이탈이나 NOL 성과는 측정하지 않는다. 대신 통합 예약의 가치가 큰 고객군을 찾아 향후 NOL 내부 데이터로 A/B 테스트할 우선순위를 제안한다.")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = footer.add_run("국내 여행로그 데이터(수도권, 2023) 기반")
set_run(fr, 8, False, GRAY)

doc.core_properties.title = "NOL 통합 예약 패키지 수요 분석 기획서"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT)
